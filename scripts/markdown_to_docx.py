#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt
from pathlib import Path
import re

md_path = Path('doc/Informe_Final.md')
out_path = Path('doc/Informe_Final.docx')

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

def strip_md_inline(text):
    # basic removal of markdown markers for bold/italic/code
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    text = re.sub(r"`(.*?)`", r"\1", text)
    return text

with md_path.open(encoding='utf-8') as f:
    for line in f:
        line = line.rstrip('\n')
        if not line.strip():
            doc.add_paragraph('')
            continue
        if line.startswith('# '):
            doc.add_heading(strip_md_inline(line[2:]), level=1)
        elif line.startswith('## '):
            doc.add_heading(strip_md_inline(line[3:]), level=2)
        elif line.startswith('### '):
            doc.add_heading(strip_md_inline(line[4:]), level=3)
        elif re.match(r'^\s*-\s', line):
            # unordered list
            text = strip_md_inline(line.lstrip(' -'))
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(text)
        elif re.match(r'^\s*\d+\.\s', line):
            text = strip_md_inline(re.sub(r'^\s*\d+\.\s', '', line))
            p = doc.add_paragraph(style='List Number')
            p.add_run(text)
        else:
            # wrap long lines
            doc.add_paragraph(strip_md_inline(line))

# Save
out_path.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(out_path))
print('WROTE', out_path)
